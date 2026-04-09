"""Generate CV in DOCX and PDF formats from parsed LinkedIn data."""
import os
import html as html_mod
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def _esc(text):
    return html_mod.escape(str(text)) if text else ""


def _set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    elm = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color_hex,
    })
    shading.append(elm)


def _remove_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            borders = tcPr.makeelement(qn("w:tcBorders"), {})
            for name in ["top", "left", "bottom", "right"]:
                el = borders.makeelement(qn(f"w:{name}"), {
                    qn("w:val"): "none", qn("w:sz"): "0",
                    qn("w:space"): "0", qn("w:color"): "auto",
                })
                borders.append(el)
            tcPr.append(borders)


def _sidebar_heading(cell, text):
    p = cell.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(79, 99, 210)
    p2 = cell.add_paragraph()
    p2.space_before = Pt(0)
    p2.space_after = Pt(4)
    r2 = p2.add_run("─" * 25)
    r2.font.size = Pt(6)
    r2.font.color.rgb = RGBColor(74, 85, 104)


def _sidebar_text(cell, text, size=8, color=RGBColor(203, 213, 224), bold=False):
    p = cell.add_paragraph()
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def _main_heading(cell, text):
    p = cell.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(26, 32, 44)
    p2 = cell.add_paragraph()
    p2.space_before = Pt(0)
    p2.space_after = Pt(6)
    r2 = p2.add_run("─" * 55)
    r2.font.size = Pt(6)
    r2.font.color.rgb = RGBColor(226, 232, 240)


def generate_docx(data, photo_path, output_path):
    """Generate a professional DOCX CV."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    sidebar = table.cell(0, 0)
    main = table.cell(0, 1)
    sidebar.width = Cm(6.5)
    main.width = Cm(13.5)
    _set_cell_shading(sidebar, "1a202c")
    sidebar.paragraphs[0].clear()
    main.paragraphs[0].clear()

    # Photo
    if photo_path and os.path.exists(photo_path):
        p_photo = sidebar.paragraphs[0]
        p_photo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_photo.add_run().add_picture(photo_path, width=Cm(3))

    # Contact
    _sidebar_heading(sidebar, "Contacto")
    if data.get("email"):
        _sidebar_text(sidebar, f"✉  {data['email']}")
    if data.get("location"):
        _sidebar_text(sidebar, f"📍 {data['location']}")
    if data.get("linkedin"):
        url = data["linkedin"].replace("https://www.", "").replace("https://", "")
        _sidebar_text(sidebar, f"🔗 {url}")
    if data.get("domain"):
        _sidebar_text(sidebar, f"🌐 {data['domain']}")

    # Skills
    if data.get("skills"):
        _sidebar_heading(sidebar, "Competências")
        for skill in data["skills"]:
            _sidebar_text(sidebar, f"▸ {skill}", 7.5)

    # Languages
    if data.get("languages"):
        _sidebar_heading(sidebar, "Idiomas")
        for lang in data["languages"]:
            _sidebar_text(sidebar, lang["name"], 8, RGBColor(226, 232, 240), bold=True)
            if lang.get("level"):
                _sidebar_text(sidebar, f"  {lang['level']}", 7, RGBColor(160, 174, 192))

    # Certifications
    if data.get("certifications"):
        _sidebar_heading(sidebar, "Certificações")
        for cert in data["certifications"]:
            _sidebar_text(sidebar, f"✓ {cert}", 7.5)

    # Main - Name
    p_name = main.paragraphs[0]
    run_name = p_name.add_run(data.get("name", ""))
    run_name.bold = True
    run_name.font.size = Pt(22)
    run_name.font.color.rgb = RGBColor(26, 32, 44)

    p_title = main.add_paragraph()
    p_title.space_before = Pt(2)
    p_title.space_after = Pt(8)
    rt = p_title.add_run(data.get("title", ""))
    rt.font.size = Pt(11)
    rt.font.color.rgb = RGBColor(113, 128, 150)

    p_line = main.add_paragraph()
    p_line.space_before = Pt(0)
    p_line.space_after = Pt(8)
    rl = p_line.add_run("━" * 55)
    rl.font.size = Pt(4)
    rl.font.color.rgb = RGBColor(79, 99, 210)

    # Summary
    if data.get("summary"):
        _main_heading(main, "Perfil")
        p = main.add_paragraph()
        p.space_after = Pt(4)
        r = p.add_run(data["summary"])
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(74, 85, 104)

    # Experience
    if data.get("experience"):
        _main_heading(main, "Experiência Profissional")
        for exp in data["experience"]:
            p = main.add_paragraph()
            p.space_before = Pt(4)
            p.space_after = Pt(0)
            rr = p.add_run(exp.get("role", ""))
            rr.bold = True
            rr.font.size = Pt(9)
            rr.font.color.rgb = RGBColor(45, 55, 72)
            if exp.get("date"):
                rs = p.add_run("  |  ")
                rs.font.size = Pt(8)
                rs.font.color.rgb = RGBColor(160, 174, 192)
                rd = p.add_run(exp["date"])
                rd.font.size = Pt(8)
                rd.font.color.rgb = RGBColor(160, 174, 192)

            p2 = main.add_paragraph()
            p2.space_before = Pt(0)
            p2.space_after = Pt(1)
            rc = p2.add_run(exp.get("company", ""))
            rc.font.size = Pt(8.5)
            rc.font.color.rgb = RGBColor(79, 99, 210)
            rc.bold = True
            if exp.get("location"):
                rloc = p2.add_run(f"  ·  {exp['location']}")
                rloc.font.size = Pt(7.5)
                rloc.font.color.rgb = RGBColor(160, 174, 192)

            if exp.get("description"):
                p3 = main.add_paragraph()
                p3.space_before = Pt(1)
                p3.space_after = Pt(4)
                rdesc = p3.add_run(exp["description"])
                rdesc.font.size = Pt(8)
                rdesc.font.color.rgb = RGBColor(74, 85, 104)

    # Education
    if data.get("education"):
        _main_heading(main, "Formação Académica")
        for edu in data["education"]:
            p = main.add_paragraph()
            p.space_before = Pt(4)
            p.space_after = Pt(0)
            rs = p.add_run(edu.get("school", ""))
            rs.bold = True
            rs.font.size = Pt(9)
            rs.font.color.rgb = RGBColor(45, 55, 72)

            p2 = main.add_paragraph()
            p2.space_before = Pt(0)
            p2.space_after = Pt(2)
            if edu.get("course"):
                rc = p2.add_run(edu["course"])
                rc.font.size = Pt(8.5)
                rc.font.color.rgb = RGBColor(79, 99, 210)
                rc.bold = True
            if edu.get("date"):
                rd = p2.add_run(f"  ·  {edu['date']}")
                rd.font.size = Pt(7.5)
                rd.font.color.rgb = RGBColor(160, 174, 192)

    _remove_table_borders(table)
    doc.save(output_path)
    return output_path


def generate_cv_html(data, photo_filename="photo.png"):
    """Generate CV HTML for PDF printing."""
    name = _esc(data.get("name", ""))
    title = _esc(data.get("title", ""))
    email = _esc(data.get("email", ""))
    location = _esc(data.get("location", ""))
    linkedin = _esc(data.get("linkedin", "").replace("https://www.", "").replace("https://", ""))
    domain = _esc(data.get("domain", ""))
    summary = _esc(data.get("summary", ""))

    # Skills
    skills_html = ""
    for s in data.get("skills", []):
        skills_html += f'<li>{_esc(s)}</li>\n'

    # Languages
    langs_html = ""
    for lang in data.get("languages", []):
        pct = {"native": 100, "bilingual": 100, "professional": 80, "limited": 50, "elementary": 30}.get("native", 50)
        level_l = (lang.get("level", "") or "").lower()
        if "native" in level_l or "bilingual" in level_l or "nativ" in level_l or "bilingue" in level_l:
            pct = 100
        elif "full professional" in level_l or "fluent" in level_l:
            pct = 90
        elif "professional" in level_l or "profissional" in level_l:
            pct = 80
        elif "limited" in level_l or "intermedi" in level_l:
            pct = 50
        elif "elementary" in level_l or "elementar" in level_l:
            pct = 30
        langs_html += f"""<div class="lang-item">
            <div class="lang-name">{_esc(lang["name"])}</div>
            <div class="lang-level">{_esc(lang.get("level", ""))}</div>
            <div class="lang-bar-container"><div class="lang-bar-fill" style="width:{pct}%"></div></div>
        </div>\n"""

    # Certifications
    certs_html = ""
    for c in data.get("certifications", []):
        certs_html += f'<div class="cert-item">{_esc(c)}</div>\n'

    # Experience
    exp_html = ""
    for exp in data.get("experience", []):
        desc = f'<div class="exp-desc">{_esc(exp.get("description", ""))}</div>' if exp.get("description") else ""
        loc = f'<div class="exp-location">{_esc(exp.get("location", ""))}</div>' if exp.get("location") else ""
        exp_html += f"""<div class="exp-item">
            <div class="exp-header"><span class="exp-role">{_esc(exp.get("role", ""))}</span><span class="exp-date">{_esc(exp.get("date", ""))}</span></div>
            <div class="exp-company">{_esc(exp.get("company", ""))}</div>
            {loc}{desc}
        </div>\n"""

    # Education
    edu_html = ""
    for edu in data.get("education", []):
        course = f'<div class="edu-course">{_esc(edu.get("course", ""))}</div>' if edu.get("course") else ""
        date = f'<div class="edu-date">{_esc(edu.get("date", ""))}</div>' if edu.get("date") else ""
        edu_html += f"""<div class="edu-item">
            <div class="edu-school">{_esc(edu.get("school", ""))}</div>
            {course}{date}
        </div>\n"""

    # Summary section
    summary_section = ""
    if summary:
        summary_section = f"""<div class="section"><h2>Perfil</h2><p class="exp-desc">{summary}</p></div>"""

    # Contact lines
    contact_lines = ""
    if email:
        contact_lines += f'<div class="contact-item"><span>✉  {email}</span></div>'
    if location:
        contact_lines += f'<div class="contact-item"><span>📍 {location}</span></div>'
    if linkedin:
        contact_lines += f'<div class="contact-item"><span>🔗 {linkedin}</span></div>'
    if domain:
        contact_lines += f'<div class="contact-item"><span>🌐 {domain}</span></div>'

    return f"""<!DOCTYPE html>
<html lang="pt">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>CV - {name}</title>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap" rel="stylesheet">
    <style>
        *{{margin:0;padding:0;box-sizing:border-box}}
        @page{{size:A4;margin:0}}
        html,body{{margin:0!important;padding:0!important;height:100%;width:100%}}
        body{{font-family:'Inter','Segoe UI',sans-serif;color:#2d3748;background:#fff;font-size:9.5pt;line-height:1.4;-webkit-print-color-adjust:exact;print-color-adjust:exact}}
        .page{{width:100%;min-height:297mm;margin:0!important;padding:0!important;display:flex;position:absolute;top:0;left:0;right:0}}
        .sidebar{{width:72mm;background:#1a202c;color:#e2e8f0;padding:24px 18px;display:flex;flex-direction:column}}
        .sidebar h2{{color:#4f63d2;font-size:9pt;text-transform:uppercase;letter-spacing:1.5px;margin-bottom:12px;padding-bottom:6px;border-bottom:1px solid #4a5568;font-weight:600}}
        .photo-container{{text-align:center;margin-bottom:20px}}
        .photo-container img{{width:90px;height:90px;border-radius:50%;object-fit:cover;object-position:center top;border:3px solid #4f63d2}}
        .sidebar-section{{margin-bottom:20px}}
        .contact-item{{display:flex;align-items:flex-start;margin-bottom:8px;font-size:8.5pt;color:#cbd5e0}}
        .skill-list{{list-style:none}}
        .skill-list li{{font-size:8.5pt;margin-bottom:5px;padding-left:12px;position:relative;color:#cbd5e0}}
        .skill-list li::before{{content:'▸';position:absolute;left:0;color:#4f63d2}}
        .lang-item{{margin-bottom:8px}}
        .lang-name{{font-size:8.5pt;font-weight:500;color:#e2e8f0;margin-bottom:3px}}
        .lang-level{{font-size:7.5pt;color:#a0aec0}}
        .lang-bar-container{{height:4px;background:#4a5568;border-radius:2px;margin-top:3px}}
        .lang-bar-fill{{height:100%;border-radius:2px;background:linear-gradient(90deg,#4f63d2,#0ea5e9)}}
        .cert-item{{font-size:8.5pt;color:#cbd5e0;padding-left:12px;position:relative;margin-bottom:4px}}
        .cert-item::before{{content:'✓';position:absolute;left:0;color:#0ea5e9;font-weight:700}}
        .main-content{{flex:1;padding:24px 22px}}
        .header{{margin-bottom:20px;border-bottom:2px solid #4f63d2;padding-bottom:14px}}
        .header h1{{font-size:22pt;font-weight:700;color:#1a202c;margin-bottom:4px}}
        .header .subtitle{{font-size:11pt;color:#718096;font-weight:400}}
        .section{{margin-bottom:18px}}
        .section h2{{font-size:11pt;font-weight:700;color:#1a202c;text-transform:uppercase;letter-spacing:1px;margin-bottom:10px;padding-bottom:4px;border-bottom:1px solid #e2e8f0}}
        .exp-item{{margin-bottom:12px;page-break-inside:avoid}}
        .exp-header{{display:flex;justify-content:space-between;align-items:baseline;margin-bottom:2px}}
        .exp-role{{font-size:9.5pt;font-weight:600;color:#2d3748}}
        .exp-date{{font-size:8pt;color:#a0aec0;white-space:nowrap}}
        .exp-company{{font-size:9pt;color:#4f63d2;font-weight:500;margin-bottom:1px}}
        .exp-location{{font-size:8pt;color:#a0aec0;margin-bottom:3px}}
        .exp-desc{{font-size:8.5pt;color:#4a5568;line-height:1.45;text-align:justify}}
        .edu-item{{margin-bottom:10px}}
        .edu-school{{font-size:9.5pt;font-weight:600;color:#2d3748}}
        .edu-course{{font-size:8.5pt;color:#4f63d2;font-weight:500}}
        .edu-date{{font-size:8pt;color:#a0aec0}}
        @media print{{body{{background:#fff}}.page{{box-shadow:none;width:100%}}}}
        @media screen{{.page{{box-shadow:0 0 20px rgba(0,0,0,.15);margin:20px auto;width:210mm}}}}
    </style>
</head>
<body>
    <div class="page">
        <aside class="sidebar">
            <div class="photo-container"><img src="{_esc(photo_filename)}" alt="{name}"></div>
            <div class="sidebar-section"><h2>Contacto</h2>{contact_lines}</div>
            <div class="sidebar-section"><h2>Competências</h2><ul class="skill-list">{skills_html}</ul></div>
            <div class="sidebar-section"><h2>Idiomas</h2>{langs_html}</div>
            {"<div class='sidebar-section'><h2>Certificações</h2>" + certs_html + "</div>" if certs_html else ""}
        </aside>
        <main class="main-content">
            <div class="header"><h1>{name}</h1><div class="subtitle">{title}</div></div>
            {summary_section}
            <div class="section"><h2>Experiência Profissional</h2>{exp_html}</div>
            <div class="section"><h2>Formação Académica</h2>{edu_html}</div>
        </main>
    </div>
</body>
</html>"""


def generate_cv_pdf(html_content, output_path):
    """Generate PDF from CV HTML using Playwright."""
    import tempfile
    from playwright.sync_api import sync_playwright

    with tempfile.NamedTemporaryFile(suffix=".html", delete=False, mode="w", encoding="utf-8") as f:
        f.write(html_content)
        tmp_path = f.name

    try:
        with sync_playwright() as p:
            browser = p.chromium.launch()
            page = browser.new_page()
            page.goto(f"file:///{tmp_path}")
            page.pdf(
                path=output_path,
                format="A4",
                print_background=True,
                margin={"top": "0", "bottom": "0", "left": "0", "right": "0"},
            )
            browser.close()
    finally:
        os.unlink(tmp_path)

    return output_path
